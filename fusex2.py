import streamlit as st
from streamlit_gsheets import GSheetsConnection
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from num2words import num2words
import os
import sys
import subprocess
import tempfile
from datetime import datetime
import pytz 
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm
from io import BytesIO
import re
import pdfplumber
import fitz  # PyMuPDF
import easyocr 
import numpy as np
import gc
import bcrypt
import time

# --- CONFIGURAÇÃO INICIAL ---
st.set_page_config(page_title="Corpore - Acesso Seguro", layout="wide", page_icon="🏥")

# --- CONEXÃO GOOGLE SHEETS ---
conn = st.connection("gsheets", type=GSheetsConnection)

# ==========================================
# 🛠 FUNÇÕES UTILITÁRIAS
# ==========================================

def limpar_data_sem_ano(texto):
    """Remove o ano de datas, mantendo dd/mm ou intervalos."""
    if pd.isna(texto): return ""
    texto = str(texto)
    texto = re.sub(r'/\d{4}', '', texto)
    texto = re.sub(r'/\d{2}(?!\d)', '', texto)
    return texto.strip()

def formatar_moeda_br(valor):
    try:
        return f"R$ {float(valor):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except:
        return "R$ 0,00"

def enviar_impressao_direta(buffer_arquivo, nome_arquivo="temp_print.docx"):
    """Salva o arquivo temporariamente e manda imprimir via terminal do SO"""
    try:
        temp_dir = tempfile.gettempdir()
        caminho_temp = os.path.join(temp_dir, nome_arquivo)
        
        # Salva o buffer em disco
        with open(caminho_temp, "wb") as f:
            f.write(buffer_arquivo.getvalue())
            
        # Comando para Windows
        if os.name == 'nt':
            os.startfile(caminho_temp, "print")
            return True, "Enviado para a impressora padrão (Windows)!"
        
        # Comando para Linux/Mac (CUPS)
        else:
            subprocess.run(["lp", caminho_temp], check=True)
            return True, "Enviado para a impressora (Linux/Mac)!"
            
    except Exception as e:
        return False, f"Erro na impressão direta: {str(e)}"

# ==========================================
# 🔐 SEGURANÇA E USUÁRIOS
# ==========================================

def carregar_usuarios():
    try:
        df = conn.read(worksheet="usuarios", ttl=0)
        return df
    except:
        return pd.DataFrame(columns=["username", "name", "password_hash", "created_at"])

def salvar_novo_usuario(username, name, password):
    df_users = carregar_usuarios()
    if not df_users.empty and username in df_users['username'].values:
        return False, "Usuário já existe!"
    
    hashed = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
    novo_usuario = pd.DataFrame([{
        "username": username, "name": name, "password_hash": hashed,
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    }])
    
    df_final = pd.concat([df_users, novo_usuario], ignore_index=True)
    try:
        conn.update(worksheet="usuarios", data=df_final)
        return True, "Usuário criado com sucesso!"
    except Exception as e:
        return False, f"Erro ao salvar: {e}"

def autenticar_usuario(username, password):
    df_users = carregar_usuarios()
    if df_users.empty: return False, "Nenhum usuário cadastrado."
    
    user_row = df_users[df_users['username'] == username]
    if user_row.empty: return False, "Usuário não encontrado."
    
    stored_hash = user_row.iloc[0]['password_hash']
    if bcrypt.checkpw(password.encode('utf-8'), stored_hash.encode('utf-8')):
        return True, user_row.iloc[0]['name']
    else:
        return False, "Senha incorreta."

def tela_login():
    st.markdown("## 🏥 Corpore Centro de Saúde")
    tab1, tab2 = st.tabs(["🔓 Entrar", "📝 Criar Conta"])
    
    with tab1:
        with st.form("login_form"):
            user = st.text_input("Usuário")
            senha = st.text_input("Senha", type="password")
            submitted = st.form_submit_button("Acessar Sistema", type="primary")
            if submitted:
                if not user or not senha: st.warning("Preencha todos os campos.")
                else:
                    sucesso, msg = autenticar_usuario(user, senha)
                    if sucesso:
                        st.session_state['logado'] = True; st.session_state['usuario_nome'] = msg
                        st.success("Login realizado!"); time.sleep(1); st.rerun()
                    else: st.error(msg)
    
    with tab2:
        with st.form("register_form"):
            new_user = st.text_input("Escolha um Usuário (Login)")
            new_name = st.text_input("Seu Nome Completo")
            new_pass = st.text_input("Escolha uma Senha", type="password")
            new_pass_conf = st.text_input("Confirme a Senha", type="password")
            reg_submitted = st.form_submit_button("Cadastrar Usuário")
            if reg_submitted:
                if new_pass != new_pass_conf: st.error("As senhas não coincidem!")
                elif len(new_pass) < 4: st.error("A senha deve ter pelo menos 4 caracteres.")
                elif not new_user: st.error("O campo usuário é obrigatório.")
                else:
                    sucesso, msg = salvar_novo_usuario(new_user, new_name, new_pass)
                    if sucesso: st.success(msg); st.info("Faça login na aba 'Entrar'.")
                    else: st.error(msg)

def logout():
    st.session_state['logado'] = False; st.session_state['usuario_nome'] = ""; st.rerun()

# ==========================================
# 🏥 SISTEMA PRINCIPAL
# ==========================================

def sistema_principal():
    with st.sidebar:
        st.write(f"Olá, **{st.session_state['usuario_nome']}**! 👋")
        if st.button("Sair / Logout"): logout()
        st.divider()

    st.title("🏥 Gestão de Faturas e Guias")

    # --- BANCO DE DADOS (SHEETS) ---
    def carregar_dados_sheets():
        try:
            df = conn.read(worksheet="guias", ttl=5)
            if not df.empty and 'fatura_ref' in df.columns:
                df['fatura_ref'] = df['fatura_ref'].astype(str).str.replace("'", "", regex=False)
            return df
        except:
            return pd.DataFrame(columns=["fatura_ref", "mes_competencia", "ano_competencia", "tipo_usuario", "servicos_fatura", "paciente_nome", "nr_guia", "prec_cp", "data_atend", "cod_proced", "valor", "data_lancamento"])

    def salvar_no_sheets(df_novo, meta_dados):
        try: df_existente = conn.read(worksheet="guias", ttl=5)
        except: df_existente = pd.DataFrame()
        
        data_hoje = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        lista_novos = []
        fatura_ref_safe = f"'{meta_dados['fatura']}"

        for _, row in df_novo.iterrows():
            try: val = float(row['VALOR (R$)'])
            except: val = 0.0
            
            prec_valor = row.get('PREC-CP/SIAPE', '')
            if pd.isna(prec_valor): prec_valor = ""

            lista_novos.append({
                "fatura_ref": fatura_ref_safe,
                "mes_competencia": meta_dados['mes'],
                "ano_competencia": meta_dados['ano'],
                "tipo_usuario": meta_dados['usuario'],
                "servicos_fatura": meta_dados['servico'],
                "paciente_nome": row['NOME DO PACIENTE'],
                "nr_guia": row['NR DA GUIA'],
                "prec_cp": str(prec_valor),
                "data_atend": limpar_data_sem_ano(row['DATA ATEND.']), 
                "cod_proced": row['CÓDIGO PROCED.'],
                "valor": val,
                "data_lancamento": data_hoje
            })
        
        df_append = pd.DataFrame(lista_novos)
        if not df_existente.empty:
            df_final = pd.concat([df_existente, df_append], ignore_index=True)
        else:
            df_final = df_append
        
        try: conn.update(worksheet="guias", data=df_final)
        except: conn.update(worksheet=0, data=df_final)

    def atualizar_fatura_sheets(fatura_ref, df_editado, meta_dados):
        df_completo = carregar_dados_sheets()
        df_limpo = df_completo[df_completo['fatura_ref'] != fatura_ref]
        
        data_hoje = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        lista_novos = []
        fatura_ref_safe = f"'{fatura_ref}"

        for _, row in df_editado.iterrows():
            try: val = float(row['VALOR (R$)'])
            except: val = 0.0
            
            prec_cp_val = row.get('PREC-CP/SIAPE', '')
            if not prec_cp_val and 'prec_cp' in row: prec_cp_val = row['prec_cp']

            lista_novos.append({
                "fatura_ref": fatura_ref_safe,
                "mes_competencia": meta_dados['mes'],
                "ano_competencia": meta_dados['ano'], 
                "tipo_usuario": meta_dados['usuario'],
                "servicos_fatura": meta_dados['servico'], 
                "paciente_nome": row['NOME DO PACIENTE'],
                "nr_guia": row['NR DA GUIA'], 
                "prec_cp": str(prec_cp_val),
                "data_atend": limpar_data_sem_ano(row['DATA ATEND.']), 
                "cod_proced": row['CÓDIGO PROCED.'], 
                "valor": val, 
                "data_lancamento": data_hoje
            })
            
        df_append = pd.DataFrame(lista_novos)
        df_limpo['fatura_ref'] = df_limpo['fatura_ref'].apply(lambda x: f"'{x}" if not str(x).startswith("'") else x)
        df_final = pd.concat([df_limpo, df_append], ignore_index=True)
        conn.update(worksheet="guias", data=df_final)

    # --- OCR / PDF ---
    @st.cache_resource
    def load_ocr_reader():
        return easyocr.Reader(['pt'], gpu=False, quantize=True)

    def extrair_texto_hibrido(arquivo_bytes):
        texto_final = ""; usou_ocr = False
        with pdfplumber.open(arquivo_bytes) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t: texto_final += t + "\n"
        
        if len(texto_final.strip()) < 20:
            try:
                reader = load_ocr_reader() 
                arquivo_bytes.seek(0)
                doc = fitz.open(stream=arquivo_bytes.read(), filetype="pdf")
                texto_ocr = ""
                for pagina in doc:
                    pix = pagina.get_pixmap(dpi=150) 
                    img_data = pix.tobytes("png")
                    resultado = reader.readtext(img_data, detail=0, paragraph=True)
                    texto_ocr += "\n".join(resultado) + "\n"
                    del pix, img_data; gc.collect()
                texto_final = texto_ocr; usou_ocr = True
            except Exception as e: return f"ERRO_OCR: {str(e)}", False
        return texto_final, usou_ocr

    def extrair_dados_pdf(arquivo):
        dados = {"NOME DO PACIENTE": "", "NR DA GUIA": "", "DATA ATEND.": "", "PREC-CP/SIAPE": "", "CÓDIGO PROCED.": "", "VALOR (R$)": 0.0}
        try:
            arquivo.seek(0)
            text, _ = extrair_texto_hibrido(arquivo)
            
            match_guia = re.search(r'(?:Nr|Numero)[:\.]?\s*(\d+)', text, flags=re.IGNORECASE)
            if match_guia: dados["NR DA GUIA"] = match_guia.group(1)
            
            match_data = re.search(r'Data:\s*(\d{2}/\d{2}/\d{4})', text, flags=re.IGNORECASE)
            if match_data: dados["DATA ATEND."] = limpar_data_sem_ano(match_data.group(1))
            else:
                datas = re.findall(r'\d{2}/\d{2}/\d{4}', text)
                if datas: dados["DATA ATEND."] = limpar_data_sem_ano(datas[0])
            
            match_titular = re.search(r'Titular:\s*\(.*?\)\s*\n?(.+)', text, flags=re.IGNORECASE)
            match_dependente = re.search(r'Dependente:\s*\(.*?\)\s*\n?(.+)', text, flags=re.IGNORECASE)
            if match_dependente: dados["NOME DO PACIENTE"] = match_dependente.group(1).strip()
            elif match_titular: dados["NOME DO PACIENTE"] = match_titular.group(1).strip()
            if "UG Origem" in dados["NOME DO PACIENTE"]: dados["NOME DO PACIENTE"] = dados["NOME DO PACIENTE"].split("UG Origem")[0].strip()
            
            match_idt = re.search(r'Idt:\s*([\d-]+)', text, flags=re.IGNORECASE)
            if match_idt: dados["PREC-CP/SIAPE"] = match_idt.group(1)
            else:
                match_prec = re.search(r'Prec CP:\s*(\d+)', text, flags=re.IGNORECASE)
                if match_prec: dados["PREC-CP/SIAPE"] = match_prec.group(1)
            
            codigos = re.findall(r'(?<!\d)(\d{8})(?!\d)', text)
            if codigos:
                codigos_validos = [c for c in codigos if not c.startswith("202")]
                dados["CÓDIGO PROCED."] = ", ".join(sorted(set(codigos_validos)))
            
            match_total = re.search(r'Total\s*:?\s*([\d\.,]+)', text, flags=re.IGNORECASE)
            if match_total:
                 try: dados["VALOR (R$)"] = float(match_total.group(1).replace('.', '').replace(',', '.'))
                 except: pass
        except: pass
        return dados

    # --- DOCX E PDF ---
    def gerar_doc_word(doc, df_dados, tags, tipo_usuario):
        for p in doc.paragraphs:
            for key, val in tags.items(): 
                if key in p.text: p.text = p.text.replace(key, str(val))
            if "REFERENTE A USUÁRIO" in p.text:
                opcoes = ["FUSEX", "PASS (S.CIVIL)", "FATOR DE CUSTO", "Ex-Combatente"]
                texto_base = "REFERENTE A USUÁRIO:   "
                for op in opcoes: 
                    marcador = "( X )" if tipo_usuario == op else "( )"
                    if tipo_usuario == "PASS" and op == "PASS (S.CIVIL)": marcador = "( X )"
                    if tipo_usuario == "S.CIVIL" and op == "PASS (S.CIVIL)": marcador = "( X )"
                    texto_base += f"{op} {marcador}    "
                p.text = texto_base
        
        if doc.tables:
            tabela = doc.tables[0]
            for _, row in df_dados.iterrows():
                cells = tabela.add_row().cells
                valor_fmt = f"{row.get('VALOR (R$)', 0.0):,.2f}".replace('.', 'X').replace(',', '.').replace('X', ',')
                vals = [row.get("NOME DO PACIENTE", ""), row.get("NR DA GUIA", ""), row.get("DATA ATEND.", ""), row.get("PREC-CP/SIAPE", ""), row.get("CÓDIGO PROCED.", ""), valor_fmt]
                for i, v in enumerate(vals): 
                    cells[i].text = str(v)
                    if cells[i].paragraphs: cells[i].paragraphs[0].runs[0].font.size = Pt(9)
            row_total = tabela.add_row().cells
            row_total[4].text = "TOTAL"; row_total[5].text = tags["{{TOTAL}}"]

        section = doc.sections[0]
        footer = section.footer
        p_footer = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        
        fuso_br = pytz.timezone('America/Sao_Paulo')
        agora = datetime.now(fuso_br).strftime("%d/%m/%Y às %H:%M")
        p_footer.text = f"Gestão Corpore - Documento gerado em: {agora}"
        p_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if not p_footer.runs: p_footer.add_run()
        p_footer.runs[0].font.size = Pt(8)
        return doc

    def criar_template_padrao():
        doc = Document()
        style = doc.styles['Normal']; style.font.name = 'Arial'; style.font.size = Pt(10)
        p = doc.add_paragraph('Corpore Centro de Saúde Ltda - CNPJ 15.259.434/0001-88')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].bold = True
        doc.add_paragraph('')
        p_fat = doc.add_paragraph(); p_fat.add_run('FATURA Nº: ').bold = True; p_fat.add_run('{{NUM_FATURA}} – {{SERVICO}} – {{MES_ANO}}')
        doc.add_paragraph('REFERENTE A USUÁRIO:   FUSEX( ) ...') 
        table = doc.add_table(rows=1, cols=6); table.style = 'Table Grid'
        hdr = ["NOME DO PACIENTE", "NR DA GUIA", "DATA ATEND.", "PREC-CP/SIAPE", "CÓD PROCED.", "VALOR R$"]
        for i, h in enumerate(hdr): table.rows[0].cells[i].text = h
        doc.add_paragraph(''); p_ext = doc.add_paragraph(); p_ext.add_run('VALOR POR EXTENSO: ').bold = True; p_ext.add_run('{{EXTENSO}} ({{TOTAL}})')
        return doc

    def gerar_pdf_protocolo(faturas_selecionadas, qtd_guias, total_faturas):
        buffer = BytesIO(); c = canvas.Canvas(buffer, pagesize=A4)
        data_envio = datetime.now().strftime("%d/%m/%Y às %H:%M")
        endereco_fusex = ["Aos Cuidados FUSEX", "Hospital Geral de Juiz de Fora - HGeJF", "Endereço: R. Gen. Deschamps Cavalcante, s/n - Fábrica", "Juiz de Fora - MG, 36080-220"]
        def desenhar_via(y_inicial):
            c.setFont("Helvetica-Bold", 14); c.drawString(2*cm, y_inicial, "CORPORE CENTRO DE SAÚDE LTDA")
            c.setFont("Helvetica", 10); c.drawString(2*cm, y_inicial - 0.5*cm, "PROTOCOLO DE REMESSA DE FATURAS FUSEX")
            c.setFont("Helvetica", 9); y_dest = y_inicial
            for linha in endereco_fusex: c.drawRightString(19*cm, y_dest, linha); y_dest -= 0.4*cm
            y_box = y_inicial - 2.8*cm; c.rect(2*cm, y_box - 2.5*cm, 17*cm, 2.5*cm)
            c.setFont("Helvetica-Bold", 11); c.drawString(2.5*cm, y_box - 0.8*cm, f"QTD FATURAS: {len(faturas_selecionadas)}"); c.drawString(10*cm, y_box - 0.8*cm, f"TOTAL DE GUIAS: {qtd_guias}")
            lista_faturas_str = [str(f) for f in faturas_selecionadas]; texto_faturas = ", ".join(lista_faturas_str)
            if len(texto_faturas) > 90: texto_faturas = texto_faturas[:90] + "..."
            c.setFont("Helvetica", 10); c.drawString(2.5*cm, y_box - 1.5*cm, f"Ref. Faturas: {texto_faturas}")
            valor_fmt = f"{total_faturas:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'); c.drawString(2.5*cm, y_box - 2.2*cm, f"Gestão Corpore: Walter Vilaça")
            y_ass = y_box - 4.5*cm; c.line(2*cm, y_ass, 9*cm, y_ass); c.setFont("Helvetica", 8); c.drawString(2*cm, y_ass - 0.4*cm, "Despachado por (Corpore)")
            c.line(11*cm, y_ass, 19*cm, y_ass); c.drawString(11*cm, y_ass - 0.4*cm, "Transportado por (Motoboy)")
            y_ass2 = y_ass - 2.5*cm; c.line(2*cm, y_ass2, 19*cm, y_ass2); c.setFont("Helvetica-Bold", 9); c.drawString(2*cm, y_ass2 - 0.5*cm, "Recebido por (Carimbo/Assinatura HGeJF)")
            c.setFont("Helvetica-Oblique", 7); c.drawRightString(19*cm, y_ass2 - 1.2*cm, f"Gerado via Sistema Corpore em: {data_envio}")
        desenhar_via(27*cm); c.setDash(4, 4); c.line(1*cm, 14.85*cm, 20*cm, 14.85*cm); c.setFont("Helvetica", 6); c.drawCentredString(10.5*cm, 14.95*cm, "- - - Corte Aqui - - -"); c.setDash([]); desenhar_via(13*cm)
        c.save(); buffer.seek(0); return buffer

    # --- INTERFACE (ABAS) ---
    tab1, tab2, tab3, tab4 = st.tabs(["📝 Nova Fatura", "✏ Editar (Nuvem)", "📊 Relatórios e 2ª Via", "📦 Protocolo"])
    meses = {"Janeiro": 1, "Fevereiro": 2, "Março": 3, "Abril": 4, "Maio": 5, "Junho": 6, "Julho": 7, "Agosto": 8, "Setembro": 9, "Outubro": 10, "Novembro": 11, "Dezembro": 12}

    # === ABA 1: NOVA FATURA ===
    with tab1:
        st.header("📝 Nova Fatura")
        if 'df_input' not in st.session_state: 
            st.session_state['df_input'] = pd.DataFrame({
                "NOME DO PACIENTE": pd.Series(dtype='str'),
                "NR DA GUIA": pd.Series(dtype='str'),
                "DATA ATEND.": pd.Series(dtype='str'),
                "PREC-CP/SIAPE": pd.Series(dtype='str'),
                "CÓDIGO PROCED.": pd.Series(dtype='str'),
                "VALOR (R$)": pd.Series(dtype='float')
            })

        c1, c2, c3 = st.columns(3)
        mes_nome = c1.selectbox("Mês", list(meses.keys()), index=datetime.now().month - 1)
        seq = c1.number_input("Sequencial", 1, 100, 1)
        fatura_ref = f"{meses[mes_nome]}.{seq}"
        c1.info(f"Fatura: **{fatura_ref}**")
        
        ano = c2.number_input("Ano", 2024, 2030, 2025)
        servico = c2.multiselect("Serviço", ["Consulta", "Fisioterapia", "Fonoaudiologia", "Psicologia", "Terapia Ocupacional", "Terapias Especiais TEA/TGD"], default=["Fisioterapia"])
        servico_txt = ", ".join(servico)
        usuario = c3.radio("Convênio", ["FUSEX", "PASS", "S.CIVIL"])
        
        uploaded = st.file_uploader("Arraste os PDFs", type="pdf", accept_multiple_files=True)
        if uploaded and st.button("Processar PDFs"):
            lista = []; bar = st.progress(0)
            for i, f in enumerate(uploaded):
                d = extrair_dados_pdf(f)
                if d["NR DA GUIA"]: lista.append(d)
                bar.progress((i+1)/len(uploaded))
            if lista:
                novo = pd.DataFrame(lista)
                novo["DATA ATEND."] = novo["DATA ATEND."].astype(str)
                st.session_state['df_input'] = pd.concat([st.session_state['df_input'], novo], ignore_index=True)
                st.success(f"{len(lista)} guias lidas!")
        
        st.session_state['df_input']['DATA ATEND.'] = st.session_state['df_input']['DATA ATEND.'].astype(str).replace('nan', '')
        st.session_state['df_input']['VALOR (R$)'] = pd.to_numeric(st.session_state['df_input']['VALOR (R$)'], errors='coerce').fillna(0.0)

        df_editor = st.data_editor(
            st.session_state['df_input'], 
            num_rows="dynamic",
            column_config={
                "VALOR (R$)": st.column_config.NumberColumn("Valor (R$)", format="R$ %.2f"),
                "DATA ATEND.": st.column_config.TextColumn("Data (dd/mm)", help="Texto livre: 05/11 ou 05/11 a 08/11")
            }
        )
        
        total = df_editor['VALOR (R$)'].sum()
        st.metric("Total", formatar_moeda_br(total))
        
        if st.button("💾 Salvar na Nuvem"):
            df_editor['DATA ATEND.'] = df_editor['DATA ATEND.'].astype(str).apply(limpar_data_sem_ano)
            meta = {'fatura': fatura_ref, 'mes': mes_nome, 'ano': ano, 'usuario': usuario, 'servico': servico_txt}
            salvar_no_sheets(df_editor, meta)
            
            doc = criar_template_padrao()
            extenso = num2words(total, lang='pt_BR', to='currency').upper()
            tags = {"{{NUM_FATURA}}": fatura_ref, "{{MES_ANO}}": f"{mes_nome}/{ano}", "{{SERVICO}}": servico_txt, "{{TOTAL}}": formatar_moeda_br(total), "{{EXTENSO}}": extenso}
            doc = gerar_doc_word(doc, df_editor, tags, usuario)
            buf = BytesIO(); doc.save(buf); buf.seek(0)
            
            st.download_button("📥 Download DOCX", buf, f"Fatura_{fatura_ref}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            st.success("Salvo com sucesso!")

    # === ABA 2: EDITAR ===
    with tab2:
        st.header("✏ Editar Faturas")
        df_nuvem = carregar_dados_sheets()
        if not df_nuvem.empty:
            faturas = df_nuvem['fatura_ref'].unique()
            sel_fat = st.selectbox("Editar fatura:", faturas)
            
            if sel_fat:
                df_filtrado = df_nuvem[df_nuvem['fatura_ref'] == sel_fat].copy()
                meta_orig = {'mes': df_filtrado.iloc[0]['mes_competencia'], 'ano': df_filtrado.iloc[0]['ano_competencia'], 'usuario': df_filtrado.iloc[0]['tipo_usuario'], 'servico': df_filtrado.iloc[0]['servicos_fatura']}
                st.info(f"Editando: {meta_orig['servico']} | {meta_orig['usuario']}")
                
                cols_possiveis = ["paciente_nome", "nr_guia", "prec_cp", "data_atend", "cod_proced", "valor"]
                cols_reais = [c for c in cols_possiveis if c in df_filtrado.columns]
                
                df_edit = df_filtrado[cols_reais].rename(columns={"paciente_nome": "NOME DO PACIENTE", "nr_guia": "NR DA GUIA", "prec_cp": "PREC-CP/SIAPE", "data_atend": "DATA ATEND.", "cod_proced": "CÓDIGO PROCED.", "valor": "VALOR (R$)"})
                
                df_edit["VALOR (R$)"] = pd.to_numeric(df_edit["VALOR (R$)"], errors='coerce').fillna(0.0)
                df_edit["DATA ATEND."] = df_edit["DATA ATEND."].astype(str).replace('nan', '')

                df_final_edit = st.data_editor(
                    df_edit, 
                    num_rows="dynamic",
                    column_config={
                        "VALOR (R$)": st.column_config.NumberColumn("Valor (R$)", format="R$ %.2f"),
                        "DATA ATEND.": st.column_config.TextColumn("Data (dd/mm)")
                    }
                )
                
                if st.button("🔄 Atualizar Fatura"):
                    df_final_edit['DATA ATEND.'] = df_final_edit['DATA ATEND.'].astype(str).apply(limpar_data_sem_ano)
                    atualizar_fatura_sheets(sel_fat, df_final_edit, meta_orig)
                    st.success("Atualizado!"); time.sleep(1); st.rerun()

    # === ABA 3: RELATÓRIOS E 2ª VIA ===
    with tab3:
        st.header("📊 Relatórios")
        df = carregar_dados_sheets()
        if not df.empty:
            df['valor'] = pd.to_numeric(df['valor'], errors='coerce')
            
            st.subheader("🖨️ Emissão de 2ª Via")
            faturas_2via = df['fatura_ref'].unique()
            col_sel, col_btn = st.columns([2, 1])
            sel_2via = col_sel.selectbox("Selecione a fatura para 2ª Via", faturas_2via)
            
            if sel_2via:
                # Reconstrói os dados
                df_fat = df[df['fatura_ref'] == sel_2via]
                meta_fat = {
                    'mes': df_fat.iloc[0]['mes_competencia'],
                    'ano': df_fat.iloc[0]['ano_competencia'],
                    'usuario': df_fat.iloc[0]['tipo_usuario'],
                    'servico': df_fat.iloc[0]['servicos_fatura']
                }
                
                df_tabela = df_fat[["paciente_nome", "nr_guia", "data_atend", "prec_cp", "cod_proced", "valor"]].copy()
                df_tabela.columns = ["NOME DO PACIENTE", "NR DA GUIA", "DATA ATEND.", "PREC-CP/SIAPE", "CÓDIGO PROCED.", "VALOR (R$)"]
                
                total_fat = df_tabela["VALOR (R$)"].sum()
                extenso_fat = num2words(total_fat, lang='pt_BR', to='currency').upper()
                tags_fat = {"{{NUM_FATURA}}": sel_2via, "{{MES_ANO}}": f"{meta_fat['mes']}/{meta_fat['ano']}", "{{SERVICO}}": meta_fat['servico'], "{{TOTAL}}": formatar_moeda_br(total_fat), "{{EXTENSO}}": extenso_fat}

                # Gera o DOCX em memória
                doc_2via = criar_template_padrao()
                doc_2via = gerar_doc_word(doc_2via, df_tabela, tags_fat, meta_fat['usuario'])
                buf_2via = BytesIO()
                doc_2via.save(buf_2via)
                buf_2via.seek(0)
                
                c_down, c_print = st.columns(2)
                c_down.download_button("📥 Baixar 2ª Via (DOCX)", buf_2via, f"2Via_{sel_2via}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                
                # RECURSO DE IMPRESSÃO DIRETA
                if c_print.button("🖨️ Imprimir Direto (Terminal)"):
                    buf_2via.seek(0) # Reseta ponteiro
                    ok, msg_imp = enviar_impressao_direta(buf_2via, f"Print_{sel_2via}.docx")
                    if ok: st.success(msg_imp)
                    else: st.error(msg_imp)
            
            st.divider()
            # Dashboard Original
            total_geral = df['valor'].sum()
            c1, c2 = st.columns(2)
            c1.metric("Faturamento Total Acumulado", formatar_moeda_br(total_geral))
            c1.metric("Total de Guias Processadas", len(df))
            st.bar_chart(df.groupby("mes_competencia")["valor"].sum())
            with st.expander("Ver Base de Dados Completa"):
                st.dataframe(df)

    # === ABA 4: PROTOCOLO ===
    with tab4:
        st.header("📦 Protocolo")
        df = carregar_dados_sheets()
        if not df.empty:
            sel = st.multiselect("Selecione Faturas:", df['fatura_ref'].unique())
            if sel:
                sub = df[df['fatura_ref'].isin(sel)]
                tot = sub['valor'].sum()
                qtd = sub['nr_guia'].nunique()
                st.info(f"Total: {formatar_moeda_br(tot)} ({qtd} guias)")
                if st.button("🖨 Baixar Protocolo"):
                    pdf = gerar_pdf_protocolo(sel, qtd, tot)
                    st.download_button("📥 PDF", pdf, "Protocolo.pdf", "application/pdf")

# ==========================================
# 🏁 MAIN
# ==========================================

if __name__ == "__main__":
    if 'logado' not in st.session_state: st.session_state['logado'] = False
    if st.session_state['logado']: sistema_principal()
    else: tela_login()

